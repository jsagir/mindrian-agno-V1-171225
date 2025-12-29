"""
Comprehensive PWS Content Ingestion

Ingests all PWS course materials to Pinecone following Mindrian best practices:
1. PWS Lectures and Worksheets (DOCX)
2. Course Material (Cohort 2024/2025, Slides, Related Material)
3. Framework-specific metadata tagging

Best Practices:
- Consistent schema: _id, content, title, category, type, source, tags, problem_type, framework_name
- Chunking with overlap for context preservation
- Deterministic IDs for idempotent upserts
- Framework-aware metadata for filtered retrieval
"""

import os
import re
import asyncio
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional
import httpx

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")

try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: PyMuPDF not installed. Install with: pip install pymupdf")

# Pinecone configuration
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
INDEX_HOST = "neo4j-knowledge-base-bc1849d.svc.aped-4627-b74a.pinecone.io"
NAMESPACE = "pws-materials"

# Base paths
PWS_LECTURES_PATH = Path("/home/jsagi/Mindrian/PWS - Lectures and worksheets created by Mindrian-20251219T001450Z-1-001/PWS - Lectures and worksheets created by Mindrian")
COURSE_MATERIAL_PATH = Path("/home/jsagi/course-material/Course Material")

# Framework metadata mapping (comprehensive)
FRAMEWORK_METADATA = {
    # From PWS Lectures
    "Jobs to be done": {
        "framework_name": "jtbd",
        "problem_type": "ill-defined",
        "tags": "framework, jtbd, jobs-to-be-done, customer-needs, hiring-criteria",
    },
    "S-Curve": {
        "framework_name": "s-curve",
        "problem_type": "ill-defined",
        "tags": "framework, s-curve, technology-lifecycle, innovation, adoption",
    },
    "Dominant designs": {
        "framework_name": "dominant-design",
        "problem_type": "ill-defined",
        "tags": "framework, dominant-design, industry-evolution, disruption, standards",
    },
    "Macro-Changes": {
        "framework_name": "macro-changes",
        "problem_type": "un-defined",
        "tags": "framework, macro-changes, trends, disruption, megatrends",
    },
    "Red Teaming": {
        "framework_name": "red-teaming",
        "problem_type": "well-defined",
        "tags": "framework, red-teaming, validation, challenge, assumptions",
    },
    "Reverse Saliant": {
        "framework_name": "reverse-salient",
        "problem_type": "un-defined",
        "tags": "framework, reverse-salient, bottleneck, innovation, constraints",
    },
    "Nested Hierarchies": {
        "framework_name": "nested-hierarchies",
        "problem_type": "ill-defined",
        "tags": "framework, nested-hierarchies, systems-thinking, abstraction",
    },
    # From Course Material
    "UN DEFINED": {
        "framework_name": "undefined-problems",
        "problem_type": "un-defined",
        "tags": "methodology, undefined, wicked-problems, exploration, 5-15-years",
    },
    "ILL-DEFINED": {
        "framework_name": "ill-defined-problems",
        "problem_type": "ill-defined",
        "tags": "methodology, ill-defined, strategy, validation, 1-5-years",
    },
    "WELL-DEFINED": {
        "framework_name": "well-defined-problems",
        "problem_type": "well-defined",
        "tags": "methodology, well-defined, execution, implementation, 0-12-months",
    },
    "WICKED": {
        "framework_name": "wicked-problems",
        "problem_type": "un-defined",
        "tags": "methodology, wicked-problems, complex, stakeholders",
    },
    "PORTFOLIO": {
        "framework_name": "portfolio-management",
        "problem_type": "all",
        "tags": "methodology, portfolio, opportunities, diversification",
    },
    "INTRODUCTION": {
        "framework_name": "pws-introduction",
        "problem_type": "all",
        "tags": "methodology, introduction, overview, pws-framework",
    },
    "MINTO": {
        "framework_name": "minto-pyramid",
        "problem_type": "all",
        "tags": "framework, minto, scqa, structured-thinking, communication",
    },
}


def extract_text_from_docx(filepath: Path) -> str:
    """Extract text from a DOCX file."""
    if not DOCX_AVAILABLE:
        return ""

    try:
        doc = Document(filepath)
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        return "\n\n".join(paragraphs)
    except Exception as e:
        print(f"Error extracting {filepath.name}: {e}")
        return ""


def extract_text_from_pdf(filepath: Path) -> str:
    """Extract text from a PDF file."""
    if not PDF_AVAILABLE:
        return ""

    try:
        doc = fitz.open(filepath)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"Error extracting {filepath.name}: {e}")
        return ""


def chunk_text(text: str, max_chunk_size: int = 2000, overlap: int = 200) -> List[str]:
    """Split text into chunks with overlap."""
    if len(text) <= max_chunk_size:
        return [text] if len(text) > 100 else []

    chunks = []
    paragraphs = text.split("\n\n")
    current_chunk = ""

    for para in paragraphs:
        if len(current_chunk) + len(para) + 2 <= max_chunk_size:
            current_chunk += ("\n\n" if current_chunk else "") + para
        else:
            if len(current_chunk) > 100:
                chunks.append(current_chunk)
            # Start new chunk with overlap
            overlap_text = current_chunk[-overlap:] if overlap > 0 and len(current_chunk) > overlap else ""
            current_chunk = overlap_text + ("\n\n" if overlap_text else "") + para

    if len(current_chunk) > 100:
        chunks.append(current_chunk)

    return chunks


def generate_id(content: str, source: str, index: int) -> str:
    """Generate deterministic chunk ID."""
    hash_input = f"{source}:{index}:{content[:100]}"
    return hashlib.md5(hash_input.encode()).hexdigest()[:12]


def detect_framework_from_content(text: str, filename: str) -> Dict[str, str]:
    """Detect framework from content and filename."""
    text_lower = text.lower()
    filename_lower = filename.lower()
    combined = text_lower + " " + filename_lower

    # Check for specific frameworks
    if "jobs to be done" in combined or "jtbd" in combined:
        return FRAMEWORK_METADATA.get("Jobs to be done", {})
    elif "s-curve" in combined or "s curve" in combined:
        return FRAMEWORK_METADATA.get("S-Curve", {})
    elif "dominant design" in combined:
        return FRAMEWORK_METADATA.get("Dominant designs", {})
    elif "macro-change" in combined or "macro change" in combined:
        return FRAMEWORK_METADATA.get("Macro-Changes", {})
    elif "red team" in combined:
        return FRAMEWORK_METADATA.get("Red Teaming", {})
    elif "reverse salient" in combined:
        return FRAMEWORK_METADATA.get("Reverse Saliant", {})
    elif "nested hierarch" in combined:
        return FRAMEWORK_METADATA.get("Nested Hierarchies", {})
    elif "un defined" in combined or "undefined" in combined:
        return FRAMEWORK_METADATA.get("UN DEFINED", {})
    elif "ill-defined" in combined or "ill defined" in combined:
        return FRAMEWORK_METADATA.get("ILL-DEFINED", {})
    elif "well-defined" in combined or "well defined" in combined:
        return FRAMEWORK_METADATA.get("WELL-DEFINED", {})
    elif "wicked" in combined:
        return FRAMEWORK_METADATA.get("WICKED", {})
    elif "portfolio" in combined:
        return FRAMEWORK_METADATA.get("PORTFOLIO", {})
    elif "minto" in combined or "scqa" in combined:
        return FRAMEWORK_METADATA.get("MINTO", {})
    elif "introduction" in combined:
        return FRAMEWORK_METADATA.get("INTRODUCTION", {})

    return {
        "framework_name": "general",
        "problem_type": "all",
        "tags": "pws, course-material",
    }


def process_file(filepath: Path, folder_context: str = "") -> List[Dict[str, Any]]:
    """Process a single file into Pinecone records."""
    # Extract text based on file type
    if filepath.suffix.lower() == ".docx":
        text = extract_text_from_docx(filepath)
    elif filepath.suffix.lower() == ".pdf":
        text = extract_text_from_pdf(filepath)
    else:
        return []

    if not text or len(text) < 100:
        return []

    chunks = chunk_text(text)
    if not chunks:
        return []

    # Detect framework from content
    meta = detect_framework_from_content(text, filepath.name)
    if not meta:
        meta = {
            "framework_name": folder_context.lower().replace(" ", "-") if folder_context else "general",
            "problem_type": "all",
            "tags": f"pws, {folder_context.lower()}" if folder_context else "pws",
        }

    # Determine document type
    filename_lower = filepath.stem.lower()
    if "lecture" in filename_lower:
        doc_type = "lecture"
        category = "Course Module"
    elif "workbook" in filename_lower or "worksheet" in filename_lower:
        doc_type = "workbook"
        category = "Course Module"
    elif "transcript" in filename_lower:
        doc_type = "transcript"
        category = "Course Module"
    elif "materials" in filename_lower or "guide" in filename_lower:
        doc_type = "guide"
        category = "Framework"
    elif "case" in filename_lower:
        doc_type = "case-study"
        category = "Case Study"
    elif "slide" in filename_lower or filepath.suffix.lower() == ".pptx":
        doc_type = "slides"
        category = "Course Module"
    elif "note" in filename_lower:
        doc_type = "notes"
        category = "Course Module"
    else:
        doc_type = "reference"
        category = "Reference"

    records = []
    for i, chunk in enumerate(chunks):
        chunk_id = generate_id(chunk, filepath.name, i)

        # Create title
        title = filepath.stem.replace("_", " ").replace("-", " ")
        # Clean up title
        title = re.sub(r'\s+', ' ', title).strip()
        if len(chunks) > 1:
            title += f" (Part {i+1})"

        # Determine source path
        try:
            source_path = str(filepath.relative_to(filepath.parents[2]))
        except (ValueError, IndexError):
            source_path = filepath.name

        records.append({
            "_id": f"pws-{meta['framework_name']}-{chunk_id}",
            "content": chunk.strip(),
            "title": title,
            "category": category,
            "type": doc_type,
            "source": source_path,
            "tags": meta["tags"],
            "problem_type": meta["problem_type"],
            "framework_name": meta["framework_name"],
        })

    return records


def process_pws_lectures() -> List[Dict[str, Any]]:
    """Process all PWS Lectures and Worksheets."""
    all_records = []

    if not PWS_LECTURES_PATH.exists():
        print(f"PWS Lectures path not found: {PWS_LECTURES_PATH}")
        return []

    print("\n=== Processing PWS Lectures and Worksheets ===")

    for folder in PWS_LECTURES_PATH.iterdir():
        if folder.is_dir():
            folder_name = folder.name
            print(f"\nProcessing folder: {folder_name}")

            for filepath in folder.glob("*.docx"):
                if "Zone.Identifier" in filepath.name:
                    continue

                print(f"  - {filepath.name}")
                records = process_file(filepath, folder_name)
                all_records.extend(records)
                print(f"    Created {len(records)} chunks")

    # Also process root-level files
    for filepath in PWS_LECTURES_PATH.glob("*.docx"):
        if "Zone.Identifier" in filepath.name:
            continue
        print(f"\nProcessing root file: {filepath.name}")
        records = process_file(filepath, "General")
        all_records.extend(records)
        print(f"  Created {len(records)} chunks")

    return all_records


def process_course_material() -> List[Dict[str, Any]]:
    """Process all Course Material."""
    all_records = []

    if not COURSE_MATERIAL_PATH.exists():
        print(f"Course Material path not found: {COURSE_MATERIAL_PATH}")
        return []

    print("\n=== Processing Course Material ===")

    # Process Cohort 2024
    cohort_2024 = COURSE_MATERIAL_PATH / "Cohort 2024"
    if cohort_2024.exists():
        print("\nProcessing Cohort 2024...")
        for filepath in cohort_2024.glob("*.docx"):
            if "Zone.Identifier" in filepath.name:
                continue
            print(f"  - {filepath.name}")
            records = process_file(filepath, "Cohort 2024")
            all_records.extend(records)
            print(f"    Created {len(records)} chunks")

    # Process Cohort 2025
    cohort_2025 = COURSE_MATERIAL_PATH / "Cohort 2025"
    if cohort_2025.exists():
        print("\nProcessing Cohort 2025...")
        for filepath in cohort_2025.rglob("*.docx"):
            if "Zone.Identifier" in filepath.name:
                continue
            print(f"  - {filepath.name}")
            records = process_file(filepath, "Cohort 2025")
            all_records.extend(records)
            print(f"    Created {len(records)} chunks")

    # Process Related Material (PDFs)
    related = COURSE_MATERIAL_PATH / "Related Material"
    if related.exists():
        print("\nProcessing Related Material...")
        for filepath in related.glob("*.pdf"):
            print(f"  - {filepath.name}")
            records = process_file(filepath, "Related Material")
            all_records.extend(records)
            print(f"    Created {len(records)} chunks")

    # Process Critical Frameworks
    frameworks = COURSE_MATERIAL_PATH / "Critical Frameworks of thinking"
    if frameworks.exists():
        print("\nProcessing Critical Frameworks...")
        for filepath in frameworks.glob("*.pdf"):
            print(f"  - {filepath.name}")
            records = process_file(filepath, "Critical Frameworks")
            all_records.extend(records)
            print(f"    Created {len(records)} chunks")

    return all_records


async def upsert_records(records: List[Dict[str, Any]], namespace: str = NAMESPACE) -> int:
    """Upsert records to Pinecone in batches."""
    if not PINECONE_API_KEY:
        raise ValueError("PINECONE_API_KEY not set")

    async with httpx.AsyncClient(timeout=60.0) as client:
        headers = {
            "Api-Key": PINECONE_API_KEY,
            "Content-Type": "application/json",
        }

        url = f"https://{INDEX_HOST}/records/namespaces/{namespace}/upsert"

        batch_size = 50
        total = 0

        for i in range(0, len(records), batch_size):
            batch = records[i:i + batch_size]

            response = await client.post(
                url,
                headers=headers,
                json={"records": batch},
            )
            response.raise_for_status()
            total += len(batch)
            print(f"  Upserted batch: {len(batch)} records (total: {total})")

        return total


def print_summary(records: List[Dict[str, Any]]):
    """Print summary of records."""
    print(f"\nTotal records created: {len(records)}")

    # By framework
    by_framework = {}
    for r in records:
        fw = r.get("framework_name", "unknown")
        by_framework[fw] = by_framework.get(fw, 0) + 1

    print("\nBy framework:")
    for fw, count in sorted(by_framework.items()):
        print(f"  {fw}: {count}")

    # By problem type
    by_type = {}
    for r in records:
        pt = r.get("problem_type", "unknown")
        by_type[pt] = by_type.get(pt, 0) + 1

    print("\nBy problem type:")
    for pt, count in sorted(by_type.items()):
        print(f"  {pt}: {count}")

    # By category
    by_cat = {}
    for r in records:
        cat = r.get("category", "unknown")
        by_cat[cat] = by_cat.get(cat, 0) + 1

    print("\nBy category:")
    for cat, count in sorted(by_cat.items()):
        print(f"  {cat}: {count}")


async def main(dry_run: bool = True, source: str = "all"):
    """Main function."""
    if not DOCX_AVAILABLE:
        print("Error: python-docx not installed")
        return

    all_records = []

    if source in ["all", "lectures"]:
        print("Processing PWS Lectures and Worksheets...")
        records = process_pws_lectures()
        all_records.extend(records)

    if source in ["all", "course"]:
        print("\nProcessing Course Material...")
        records = process_course_material()
        all_records.extend(records)

    if not all_records:
        print("\nNo records created. Check the input files.")
        return

    print_summary(all_records)

    if dry_run:
        print("\n[DRY RUN] Would upsert the following records:")
        for r in all_records[:15]:
            print(f"  - {r['_id']}: {r['title'][:60]}...")
        if len(all_records) > 15:
            print(f"  ... and {len(all_records) - 15} more")
        print("\nRun with --upsert to actually upload to Pinecone")
    else:
        print(f"\nUpserting to Pinecone namespace: {NAMESPACE}")
        count = await upsert_records(all_records)
        print(f"\nSuccessfully upserted {count} records")


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Ingest all PWS Content to Pinecone")
    parser.add_argument("--upsert", action="store_true", help="Actually upsert to Pinecone")
    parser.add_argument("--source", choices=["all", "lectures", "course"], default="all",
                       help="Which source to process (default: all)")
    args = parser.parse_args()

    asyncio.run(main(dry_run=not args.upsert, source=args.source))
