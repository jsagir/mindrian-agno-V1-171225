"""
Ingest PWS Lectures and Worksheets to Pinecone

Extracts text from DOCX files and ingests to PWS Brain.
"""

import os
import re
import asyncio
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional
import httpx

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")

# Pinecone configuration
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
INDEX_HOST = "neo4j-knowledge-base-bc1849d.svc.aped-4627-b74a.pinecone.io"
NAMESPACE = "pws-materials"

# Base path for lectures
BASE_PATH = Path("/home/jsagi/Mindrian/PWS - Lectures and worksheets created by Mindrian-20251219T001450Z-1-001/PWS - Lectures and worksheets created by Mindrian")

# Framework metadata mapping
FRAMEWORK_METADATA = {
    "Jobs to be done": {
        "framework_name": "jtbd",
        "problem_type": "ill-defined",
        "tags": "framework, jtbd, jobs-to-be-done, customer-needs",
    },
    "S-Curve": {
        "framework_name": "s-curve",
        "problem_type": "ill-defined",
        "tags": "framework, s-curve, technology-lifecycle, innovation",
    },
    "Dominant designs": {
        "framework_name": "dominant-design",
        "problem_type": "ill-defined",
        "tags": "framework, dominant-design, industry-evolution, disruption",
    },
    "Macro-Changes": {
        "framework_name": "macro-changes",
        "problem_type": "un-defined",
        "tags": "framework, macro-changes, trends, disruption",
    },
    "Red Teaming": {
        "framework_name": "red-teaming",
        "problem_type": "well-defined",
        "tags": "framework, red-teaming, validation, challenge",
    },
    "Reverse Saliant": {
        "framework_name": "reverse-salient",
        "problem_type": "un-defined",
        "tags": "framework, reverse-salient, bottleneck, innovation",
    },
    "Nested Hierarchies": {
        "framework_name": "nested-hierarchies",
        "problem_type": "ill-defined",
        "tags": "framework, nested-hierarchies, systems-thinking",
    },
}


def extract_text_from_docx(filepath: Path) -> str:
    """Extract text from a DOCX file."""
    if not DOCX_AVAILABLE:
        return ""

    try:
        doc = Document(filepath)
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        return "\n\n".join(paragraphs)
    except Exception as e:
        print(f"Error extracting {filepath.name}: {e}")
        return ""


def chunk_text(text: str, max_chunk_size: int = 2000, overlap: int = 200) -> List[str]:
    """Split text into chunks with overlap."""
    if len(text) <= max_chunk_size:
        return [text] if len(text) > 100 else []

    chunks = []
    paragraphs = text.split("\n\n")
    current_chunk = ""

    for para in paragraphs:
        if len(current_chunk) + len(para) + 2 <= max_chunk_size:
            current_chunk += ("\n\n" if current_chunk else "") + para
        else:
            if len(current_chunk) > 100:
                chunks.append(current_chunk)
            # Start new chunk with overlap
            overlap_text = current_chunk[-overlap:] if overlap > 0 and len(current_chunk) > overlap else ""
            current_chunk = overlap_text + ("\n\n" if overlap_text else "") + para

    if len(current_chunk) > 100:
        chunks.append(current_chunk)

    return chunks


def generate_id(content: str, source: str, index: int) -> str:
    """Generate deterministic chunk ID."""
    hash_input = f"{source}:{index}:{content[:100]}"
    return hashlib.md5(hash_input.encode()).hexdigest()[:12]


def process_docx_file(filepath: Path, folder_name: str) -> List[Dict[str, Any]]:
    """Process a single DOCX file into Pinecone records."""
    text = extract_text_from_docx(filepath)
    if not text:
        return []

    chunks = chunk_text(text)
    if not chunks:
        return []

    # Get metadata for this framework
    meta = FRAMEWORK_METADATA.get(folder_name, {
        "framework_name": folder_name.lower().replace(" ", "-"),
        "problem_type": "all",
        "tags": f"framework, {folder_name.lower()}",
    })

    # Determine document type
    filename_lower = filepath.stem.lower()
    if "lecture" in filename_lower:
        doc_type = "lecture"
    elif "workbook" in filename_lower or "worksheet" in filename_lower:
        doc_type = "workbook"
    elif "materials" in filename_lower or "guide" in filename_lower:
        doc_type = "guide"
    elif "case" in filename_lower:
        doc_type = "case-study"
    else:
        doc_type = "reference"

    records = []
    for i, chunk in enumerate(chunks):
        chunk_id = generate_id(chunk, filepath.name, i)

        # Create title
        title = filepath.stem.replace("_", " ").replace("-", " ")
        if len(chunks) > 1:
            title += f" (Part {i+1})"

        records.append({
            "_id": f"pws-{meta['framework_name']}-{chunk_id}",
            "content": chunk.strip(),
            "title": title,
            "category": "Course Module" if doc_type in ["lecture", "workbook"] else "Framework",
            "type": doc_type,
            "source": f"pws-lectures/{folder_name}/{filepath.name}",
            "tags": meta["tags"],
            "problem_type": meta["problem_type"],
            "framework_name": meta["framework_name"],
        })

    return records


def process_all_files() -> List[Dict[str, Any]]:
    """Process all DOCX files in the PWS lectures directory."""
    all_records = []

    for folder in BASE_PATH.iterdir():
        if folder.is_dir():
            folder_name = folder.name
            print(f"\nProcessing folder: {folder_name}")

            for filepath in folder.glob("*.docx"):
                # Skip Zone.Identifier files
                if "Zone.Identifier" in filepath.name:
                    continue

                print(f"  - {filepath.name}")
                records = process_docx_file(filepath, folder_name)
                all_records.extend(records)
                print(f"    Created {len(records)} chunks")

    # Also process root-level files
    for filepath in BASE_PATH.glob("*.docx"):
        if "Zone.Identifier" in filepath.name:
            continue
        print(f"\nProcessing root file: {filepath.name}")
        records = process_docx_file(filepath, "General")
        all_records.extend(records)
        print(f"  Created {len(records)} chunks")

    return all_records


async def upsert_records(records: List[Dict[str, Any]], namespace: str = NAMESPACE) -> int:
    """Upsert records to Pinecone in batches."""
    if not PINECONE_API_KEY:
        raise ValueError("PINECONE_API_KEY not set")

    async with httpx.AsyncClient(timeout=60.0) as client:
        headers = {
            "Api-Key": PINECONE_API_KEY,
            "Content-Type": "application/json",
        }

        url = f"https://{INDEX_HOST}/records/namespaces/{namespace}/upsert"

        batch_size = 50
        total = 0

        for i in range(0, len(records), batch_size):
            batch = records[i:i + batch_size]

            response = await client.post(
                url,
                headers=headers,
                json={"records": batch},
            )
            response.raise_for_status()
            total += len(batch)
            print(f"  Upserted batch: {len(batch)} records (total: {total})")

        return total


async def main(dry_run: bool = True):
    """Main function."""
    if not DOCX_AVAILABLE:
        print("Error: python-docx not installed")
        return

    print("Processing PWS Lectures and Worksheets...")
    records = process_all_files()

    if not records:
        print("\nNo records created. Check the input files.")
        return

    print(f"\nTotal records created: {len(records)}")

    # Show summary by framework
    by_framework = {}
    for r in records:
        fw = r.get("framework_name", "unknown")
        by_framework[fw] = by_framework.get(fw, 0) + 1

    print("\nBy framework:")
    for fw, count in sorted(by_framework.items()):
        print(f"  {fw}: {count}")

    if dry_run:
        print("\n[DRY RUN] Would upsert the following records:")
        for r in records[:10]:
            print(f"  - {r['_id']}: {r['title'][:50]}...")
        if len(records) > 10:
            print(f"  ... and {len(records) - 10} more")
        print("\nRun with --upsert to actually upload to Pinecone")
    else:
        print(f"\nUpserting to Pinecone namespace: {namespace}")
        count = await upsert_records(records)
        print(f"\nSuccessfully upserted {count} records")


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Ingest PWS Lectures to Pinecone")
    parser.add_argument("--upsert", action="store_true", help="Actually upsert to Pinecone")
    args = parser.parse_args()

    asyncio.run(main(dry_run=not args.upsert))
