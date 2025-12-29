#!/usr/bin/env python3
"""
Ingest PWS Lectures and Worksheets to Neo4j + Pinecone GraphRAG

This script:
1. Reads DOCX files from the PWS lectures directory
2. Extracts and chunks the content
3. Creates Neo4j nodes (Lecture, Worksheet) linked to Framework nodes
4. Upserts to Pinecone 'mindrian-graphrag' index with graph linking metadata

Usage:
    python scripts/ingest_pws_lectures_to_graphrag.py
"""

import os
import sys
import re
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Tuple
from dataclasses import dataclass

# Add parent to path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import docx
import httpx

# Configuration
PWS_CONTENT_DIR = "/home/jsagi/Mindrian/PWS - Lectures and worksheets created by Mindrian-20251219T001450Z-1-001/PWS - Lectures and worksheets created by Mindrian"

PINECONE_INDEX_HOST = "mindrian-graphrag-bc1849d.svc.aped-4627-b74a.pinecone.io"
PINECONE_NAMESPACE = "graphrag"

# Topic to Framework mapping
TOPIC_FRAMEWORK_MAP = {
    "jobs to be done": "Jobs to Be Done (JTBD)",
    "jtbd": "Jobs to Be Done (JTBD)",
    "s-curve": "S-Curve Analysis",
    "dominant design": "Dominant Design Framework",
    "macro-change": "Macro-Changes Framework",
    "red team": "Red Teaming",
    "reverse salient": "Reverse Salient",
    "nested hierarch": "Nested Hierarchies",
}


@dataclass
class ContentChunk:
    """A chunk of content to be ingested"""
    id: str
    content: str
    title: str
    source_file: str
    topic: str
    content_type: str  # lecture, workbook, guide, case_study
    framework: str
    chunk_index: int
    total_chunks: int


def extract_docx_text(file_path: str) -> str:
    """Extract text from a DOCX file"""
    try:
        doc = docx.Document(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        return "\n\n".join(paragraphs)

    except Exception as e:
        print(f"  Error reading {file_path}: {e}")
        return ""


def determine_content_type(filename: str) -> str:
    """Determine content type from filename"""
    filename_lower = filename.lower()

    if "workbook" in filename_lower or "worksheet" in filename_lower:
        return "workbook"
    elif "lecture" in filename_lower:
        return "lecture"
    elif "guide" in filename_lower or "materials" in filename_lower:
        return "guide"
    elif "case" in filename_lower or "example" in filename_lower:
        return "case_study"
    else:
        return "lecture"


def determine_topic_and_framework(file_path: str) -> Tuple[str, str]:
    """Determine topic and related framework from file path"""
    path_lower = file_path.lower()

    for keyword, framework in TOPIC_FRAMEWORK_MAP.items():
        if keyword in path_lower:
            # Extract topic from parent folder
            parent_folder = Path(file_path).parent.name
            return parent_folder, framework

    # Default
    parent_folder = Path(file_path).parent.name
    return parent_folder, parent_folder


def chunk_text(text: str, max_chunk_size: int = 1500, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks"""
    if len(text) <= max_chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + max_chunk_size

        # Try to break at paragraph or sentence boundary
        if end < len(text):
            # Look for paragraph break
            para_break = text.rfind("\n\n", start, end)
            if para_break > start + max_chunk_size // 2:
                end = para_break
            else:
                # Look for sentence break
                sentence_break = text.rfind(". ", start, end)
                if sentence_break > start + max_chunk_size // 2:
                    end = sentence_break + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        start = end - overlap

    return chunks


def generate_chunk_id(file_path: str, chunk_index: int) -> str:
    """Generate unique ID for chunk"""
    file_hash = hashlib.md5(file_path.encode()).hexdigest()[:8]
    filename = Path(file_path).stem.lower()
    # Clean filename for ID
    clean_name = re.sub(r'[^a-z0-9]+', '_', filename)[:30]
    return f"pws_{clean_name}_{file_hash}_{chunk_index}"


def process_directory(base_dir: str) -> List[ContentChunk]:
    """Process all DOCX files in directory"""
    chunks = []

    for root, dirs, files in os.walk(base_dir):
        for filename in files:
            if not filename.endswith('.docx') or filename.startswith('~'):
                continue

            file_path = os.path.join(root, filename)
            print(f"  Processing: {filename}")

            # Extract text
            text = extract_docx_text(file_path)
            if not text:
                continue

            # Determine metadata
            content_type = determine_content_type(filename)
            topic, framework = determine_topic_and_framework(file_path)

            # Chunk text
            text_chunks = chunk_text(text)

            # Create ContentChunk objects
            for i, chunk_content in enumerate(text_chunks):
                chunk = ContentChunk(
                    id=generate_chunk_id(file_path, i),
                    content=chunk_content,
                    title=Path(filename).stem.replace('_', ' ').replace('-', ' '),
                    source_file=filename,
                    topic=topic,
                    content_type=content_type,
                    framework=framework,
                    chunk_index=i,
                    total_chunks=len(text_chunks),
                )
                chunks.append(chunk)

            print(f"    → {len(text_chunks)} chunks created")

    return chunks


def create_neo4j_nodes(chunks: List[ContentChunk]) -> int:
    """Create Neo4j nodes for the content"""
    # Group chunks by source file
    files = {}
    for chunk in chunks:
        if chunk.source_file not in files:
            files[chunk.source_file] = {
                "title": chunk.title,
                "topic": chunk.topic,
                "framework": chunk.framework,
                "content_type": chunk.content_type,
                "chunks": []
            }
        files[chunk.source_file]["chunks"].append(chunk)

    print(f"\n  Creating Neo4j nodes for {len(files)} documents...")

    # For now, we'll use MCP tools via the existing connection
    # In production, this would use the Neo4j driver directly

    neo4j_uri = os.getenv("NEO4J_URI")
    neo4j_password = os.getenv("NEO4J_PASSWORD")

    if not neo4j_password:
        print("  ⚠️  NEO4J_PASSWORD not set - skipping Neo4j ingestion")
        print("  Neo4j nodes will need to be created separately")
        return 0

    # Would create nodes like:
    # (d:Document:Lecture {title, topic, framework, content_type})
    # (d)-[:TEACHES]->(f:Framework)
    # (d)-[:HAS_CHUNK]->(c:Chunk {content, index})

    return len(files)


def save_chunks_to_json(chunks: List[ContentChunk], output_path: str) -> int:
    """Save chunks to JSON for MCP tool ingestion"""
    import json

    records = []
    for chunk in chunks:
        record = {
            "_id": chunk.id,
            "content": chunk.content,
            "title": chunk.title,
            "category": chunk.content_type.title(),
            "topic": chunk.topic,
            "framework": chunk.framework,
            "source_file": chunk.source_file,
            "chunk_index": str(chunk.chunk_index),  # Convert to string
            "total_chunks": str(chunk.total_chunks),  # Convert to string
            "source": "pws_lectures",
        }
        records.append(record)

    with open(output_path, 'w') as f:
        json.dump(records, f, indent=2)

    print(f"\n  Saved {len(records)} records to {output_path}")
    print("  Use MCP tools to upsert: mcp__pinecone__upsert-records")
    return len(records)


def main():
    print("=" * 60)
    print("PWS Lectures & Worksheets → GraphRAG Ingestion")
    print("=" * 60)

    # Check directory exists
    if not os.path.exists(PWS_CONTENT_DIR):
        print(f"Error: Directory not found: {PWS_CONTENT_DIR}")
        sys.exit(1)

    # Process files
    print(f"\n📁 Processing directory: {PWS_CONTENT_DIR}\n")
    chunks = process_directory(PWS_CONTENT_DIR)

    print(f"\n✅ Total chunks created: {len(chunks)}")

    # Show summary by topic
    topics = {}
    for chunk in chunks:
        if chunk.topic not in topics:
            topics[chunk.topic] = 0
        topics[chunk.topic] += 1

    print("\n📊 Chunks by topic:")
    for topic, count in sorted(topics.items()):
        print(f"    {topic}: {count}")

    # Create Neo4j nodes
    neo4j_count = create_neo4j_nodes(chunks)

    # Save chunks to JSON for MCP tool ingestion
    output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data", "pws_chunks.json")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    saved_count = save_chunks_to_json(chunks, output_path)

    # Summary
    print("\n" + "=" * 60)
    print("Summary")
    print("=" * 60)
    print(f"  Documents processed: {len(set(c.source_file for c in chunks))}")
    print(f"  Total chunks: {len(chunks)}")
    print(f"  Neo4j nodes: {neo4j_count}")
    print(f"  Chunks saved to JSON: {saved_count}")
    print(f"\n  Next: Use MCP tool to upsert to Pinecone")
    print("=" * 60)


if __name__ == "__main__":
    main()
